from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xDC, 0x26, 0x26)
DARK = RGBColor(0x18, 0x18, 0x1B)
GREY = RGBColor(0x52, 0x52, 0x5B)

CATEGORIES = {
    "elite": {
        "label": "Elite Private Training",
        "tagline": "Private Self Defense, Personal Protection & Tactical Training",
        "description": "Discreet, founder-led private instruction for executives, professionals, families, and serious individuals who want elite-level personal protection, situational readiness, and real combat skill, taught one-on-one by Grandmaster Dr. David Furie.",
    },
    "sport": {
        "label": "Martial Arts & Combat Sports",
        "tagline": "Competitive Martial Arts & Sport Combat Training",
        "description": "Private competitive martial arts instruction in Brooklyn for adults and young athletes who want the sport side of the system, including MMA, Jujitsu, Kickboxing, Ninjutsu, Self Defense, and Weapons and Tactics.",
    },
}

SERVICES = [
    {
        "title": "Private Instruction",
        "route": "/private-instruction",
        "price": "$250 / session",
        "category": "elite",
        "subheadline": "One-on-one training tailored to your goals, experience, and preferred areas of focus.",
        "overview": "Private Instruction is designed for clients who want direct access to high-level founder-led training without the distractions of a group class. Sessions are tailored around the client's background, goals, comfort level, and priorities, with a focus on practical skill, personal readiness, and real-world application.",
        "whoFor": [
            "Serious adults seeking personalized training",
            "Professionals looking for private, high-level instruction",
            "Beginners who want to learn in a private setting",
            "Returning martial artists looking to refine their skills",
            "Clients who value discretion and individual attention",
            "Individuals who want training built around their specific goals",
        ],
        "whatYouGet": [
            "One-on-one instruction with Grandmaster Dr. David Furie",
            "Customized session focus based on your goals",
            "Direct feedback and correction throughout the session",
            "Practical skill development for real-world application",
            "A private training experience shaped around your needs",
        ],
        "trainingFocus": [
            "Movement and positioning", "Striking techniques", "Grappling fundamentals",
            "Defensive positioning", "Situational awareness", "Tactical response",
            "Pressure-point application", "Controlled performance under stress",
        ],
        "likelyOutcomes": [
            "Better confidence under pressure",
            "More awareness and control in everyday situations",
            "Improved movement, positioning, and reflexes",
            "Stronger personal readiness and self-command",
            "A clearer sense of what works for you in real life",
        ],
        "whyChoose": "This is not generic class-based training. It is direct, personalized, founder-led instruction for people who want a more serious and practical learning experience.",
        "faqs": [
            ("Do I need prior experience?", "No. Sessions are tailored to your background and comfort level, whether you are a complete beginner or an experienced practitioner."),
            ("Is this suitable for beginners?", "Absolutely. Private Instruction is one of the best ways to learn because every session is shaped around your specific needs and pace."),
            ("How are sessions customized?", "David assesses your goals, background, and areas of interest at the start and adjusts the focus, intensity, and content accordingly."),
            ("Is this offered privately only?", "Yes. All Fury Combat instruction is private. There are no group classes."),
            ("Can this focus on a specific area of concern?", "Yes. Sessions can be tailored around a specific skill, scenario, or area of readiness you want to develop."),
        ],
        "videos": [],
    },
    {
        "title": "Advanced Tactical Instruction",
        "route": "/advanced-tactical-instruction",
        "price": "$500 / session",
        "category": "elite",
        "subheadline": "An elevated private offering for serious individuals seeking instruction in readiness, tactical thinking, and strategic response.",
        "overview": "Advanced Tactical Instruction is built for clients who want a deeper and more sophisticated level of private training. It is designed around readiness, awareness under pressure, protective movement, strategic thinking, and practical response in more demanding situations.",
        "whoFor": [
            "Experienced trainees seeking elevated instruction",
            "Serious private clients with specific readiness goals",
            "Executives and high-profile individuals",
            "Security-minded professionals",
            "Individuals who want a higher level of preparedness",
        ],
        "whatYouGet": [
            "Advanced private instruction with founder-level coaching",
            "Higher-level correction and strategic feedback",
            "Scenario-oriented training emphasis",
            "More strategic and readiness-focused development",
            "Personalized sessions shaped around your concerns and goals",
        ],
        "trainingFocus": [
            "Readiness under pressure", "Tactical thinking and decision-making",
            "Awareness under pressure", "Protective movement", "Weapons familiarity",
            "Reconnaissance concepts", "Strategic response",
        ],
        "likelyOutcomes": [
            "Stronger composure in uncertain situations",
            "Better decision-making under pressure",
            "Improved awareness and readiness",
            "Deeper tactical understanding",
            "More confidence in complex environments",
        ],
        "whyChoose": "This offering is for clients who know ordinary training is not enough for what they want. It is more refined, more strategic, and more serious than standard instruction.",
        "faqs": [
            ("Is this only for military or law enforcement?", "No. This training is available to serious civilian clients who want a higher standard of readiness and tactical awareness."),
            ("Can civilians do this training?", "Yes. Many clients are professionals, executives, or individuals who simply want more advanced and strategic private instruction."),
            ("Do I need prior experience?", "Some prior training or a strong fitness base is helpful, but David will assess your level and adjust the session accordingly."),
            ("How is this different from Private Instruction?", "Advanced Tactical Instruction is more strategic, more readiness-focused, and designed for clients seeking a deeper level of training beyond standard private sessions."),
            ("Can the focus be customized?", "Yes. Every session is tailored to the client's goals, concerns, and experience level."),
        ],
        "videos": [],
    },
    {
        "title": "Women's Private Safety Training",
        "route": "/womens-private-safety-training",
        "price": "$300 / session",
        "category": "elite",
        "subheadline": "Private instruction designed to help women strengthen awareness, confidence, prevention skills, and decisive real-world response.",
        "overview": "Women's Private Safety Training is designed to help women feel more aware, more prepared, and more confident in everyday situations. The focus is not fear. The focus is practical skill, better decisions, and stronger personal presence under pressure.",
        "whoFor": [
            "Women seeking private, personalized training",
            "Professionals who commute or travel regularly",
            "Women living or working in the city",
            "Mothers and daughters seeking shared training",
            "Students and young professionals",
        ],
        "whatYouGet": [
            "Private one-on-one or small private session instruction",
            "Personalized attention and pacing",
            "Practical prevention and awareness coaching",
            "Training tailored to your comfort level and concerns",
            "A discreet and supportive learning environment",
        ],
        "trainingFocus": [
            "Awareness and environmental reading", "Confidence and personal presence",
            "Prevention skills and early recognition", "De-escalation techniques",
            "Escape options and decisive action", "Boundary setting",
            "Voice and presence under pressure",
        ],
        "likelyOutcomes": [
            "More confidence moving through public spaces",
            "Stronger awareness habits in everyday life",
            "Better boundary-setting ability",
            "Improved emotional control under stress",
            "Greater readiness to act earlier and more decisively",
        ],
        "whyChoose": "This is practical women's safety training in a private setting, not a generic seminar or fear-based class. It is personalized, direct, and built around real-world application.",
        "faqs": [
            ("Is this beginner-friendly?", "Yes. This training is designed to meet you where you are, regardless of prior experience."),
            ("Can this be tailored to a specific concern?", "Absolutely. Sessions can focus on commuting safety, travel, workplace awareness, or any personal concern."),
            ("Is it private only?", "Yes. All sessions are private, though small groups (such as mothers and daughters) can be arranged."),
            ("Is this fitness-based or practical training?", "The focus is practical awareness and safety skill, not fitness. Physical conditioning may be incorporated when relevant."),
            ("Can mothers and daughters train together?", "Yes. Private sessions can include family members when appropriate."),
        ],
        "videos": [],
    },
    {
        "title": "Tactical Conditioning",
        "route": "/tactical-conditioning",
        "price": "$150 / session",
        "category": "elite",
        "subheadline": "A private training experience that combines conditioning, coordination, movement, reaction, and practical defensive drills.",
        "overview": "Tactical Conditioning is for clients who want to improve the physical side of readiness while staying connected to practical movement and defensive application. It blends conditioning with reaction, control, and purposeful training.",
        "whoFor": [
            "Adults who want practical, purposeful conditioning",
            "Clients building physical readiness alongside tactical skill",
            "Professionals who want focused, efficient training",
            "Individuals seeking a lower barrier entry into private instruction",
            "People looking to sharpen movement and coordination",
        ],
        "whatYouGet": [
            "Private guided training with David Furie",
            "Conditioning with real-world relevance",
            "Movement and reaction drills",
            "Practical defensive applications built into the work",
            "Training scaled to your level and pace",
        ],
        "trainingFocus": [
            "Conditioning and endurance", "Coordination and body control",
            "Movement and footwork", "Reaction time and reflexes",
            "Practical defensive drills", "Balance and recovery under pressure",
        ],
        "likelyOutcomes": [
            "Improved physical readiness and stamina",
            "Sharper reactions and quicker response time",
            "Better coordination and body control",
            "Stronger endurance and movement efficiency",
            "More confidence in your physical capabilities",
        ],
        "whyChoose": "This is more purposeful than generic fitness training because it connects physical development to practical readiness and controlled application.",
        "faqs": [
            ("Is this more fitness or self-defense?", "It bridges both. The conditioning work is purposeful and connected to practical defensive movement and readiness."),
            ("Is it beginner-friendly?", "Yes. Training is scaled to your fitness level and experience."),
            ("Can this be combined with Private Instruction?", "Yes. Many clients complement their tactical conditioning with private instruction sessions."),
            ("Is it appropriate for older adults?", "Yes. Sessions are tailored to your physical condition and adjusted accordingly."),
            ("How intense is a session?", "Intensity is scaled to your level. David will assess your baseline and adjust throughout the session."),
        ],
        "videos": [],
    },
    {
        "title": "Young Adult Readiness Training",
        "route": "/young-adult-readiness-training",
        "price": "$225 / session",
        "category": "elite",
        "subheadline": "Private instruction for young adults preparing for college, commuting, travel, city life, or greater independence.",
        "overview": "Young Adult Readiness Training is designed to help young adults become more aware, more confident, and better prepared for the realities of independence. It is especially relevant for city life, commuting, travel, and transitional life stages.",
        "whoFor": [
            "High school seniors preparing for the next chapter",
            "College students navigating campus and city life",
            "Young adults entering the workforce",
            "Commuters seeking stronger awareness habits",
            "Parents seeking private training for their children",
        ],
        "whatYouGet": [
            "Private training tailored to the individual",
            "Real-world readiness coaching",
            "Awareness and decision-making support",
            "Practical skills for public-space and day-to-day situations",
            "A serious but age-appropriate learning environment",
        ],
        "trainingFocus": [
            "Commuting awareness and city safety", "Travel readiness",
            "Confidence and personal presence", "Verbal boundaries and de-escalation",
            "Practical response habits", "Situational awareness in public spaces",
        ],
        "likelyOutcomes": [
            "Stronger confidence in unfamiliar environments",
            "Better everyday awareness and observation habits",
            "Improved decision-making in real-time situations",
            "Greater readiness for independence and self-reliance",
            "More peace of mind for both the student and family",
        ],
        "whyChoose": "This helps bridge the gap between dependence and independence with practical private instruction that is serious, modern, and relevant to real life.",
        "faqs": [
            ("Is this appropriate for teens or only college-age students?", "This training is appropriate for older teens through young adulthood, typically ages 16 and up."),
            ("Can parents be involved?", "Yes. Parents are welcome to discuss goals and concerns, and can observe or participate if appropriate."),
            ("Is it private only?", "Yes. All sessions are private and tailored to the individual."),
            ("Can this focus on commuting or city life?", "Absolutely. Sessions can be tailored to the specific situations and environments the young adult will encounter."),
            ("Is prior experience required?", "No. This training is designed to meet the student where they are."),
        ],
        "videos": [],
    },
    {
        "title": "Executive Readiness",
        "route": "/executive-readiness",
        "price": "$400 / session",
        "category": "elite",
        "subheadline": "A premium private training offering for executives, entrepreneurs, professionals, and public-facing individuals who value preparedness, discretion, awareness, and self-command.",
        "overview": "Executive Readiness is built for higher-level clients who want a discreet, serious, and premium private training experience. The emphasis is on composure, awareness, preparedness, and real-world self-command.",
        "whoFor": [
            "Executives and C-suite professionals",
            "Entrepreneurs and business owners",
            "Public-facing individuals and thought leaders",
            "Clients who value discretion and privacy",
            "Higher-income private clients seeking serious instruction",
        ],
        "whatYouGet": [
            "Founder-led private training with David Furie",
            "Premium personalized instruction",
            "Discreet, focused sessions",
            "Practical readiness coaching",
            "A higher-standard training experience",
        ],
        "trainingFocus": [
            "Preparedness and personal security awareness", "Discretion and situational control",
            "Awareness in professional environments", "Self-command and composure under pressure",
            "Movement and positioning", "Decision-making under pressure",
        ],
        "likelyOutcomes": [
            "Greater confidence and composure in high-stakes situations",
            "Improved awareness in public and professional environments",
            "Stronger self-command under pressure",
            "A more personalized readiness skill set",
            "A training experience aligned with a demanding lifestyle",
        ],
        "whyChoose": "This is designed for people who do not want a public class environment and who value serious, tailored, private instruction delivered at a higher standard.",
        "faqs": [
            ("Is this only for executives?", "No. Executive Readiness is for any serious professional or individual who values privacy, quality, and a premium training experience."),
            ("How private is the training?", "Completely private. Sessions are one-on-one with David in a private setting."),
            ("Can it be customized to my work and lifestyle?", "Yes. Sessions are built around your specific schedule, concerns, and goals."),
            ("Is it beginner-friendly?", "Yes. David tailors every session to the client's current level and adjusts accordingly."),
            ("How is this different from Private Instruction?", "Executive Readiness is positioned at a higher standard with a focus on composure, discretion, and readiness for professionals in demanding roles."),
        ],
        "videos": [],
    },
    {
        "title": "Family Protection Session",
        "route": "/family-protection-session",
        "price": "$350 / session",
        "category": "elite",
        "subheadline": "A private session for individuals or families focused on practical awareness, protective habits, emergency thinking, and everyday readiness.",
        "overview": "Family Protection Sessions are designed for individuals or families who want practical guidance around awareness, household and daily-life readiness, emergency thinking, and protective habits that make real-world situations more manageable.",
        "whoFor": [
            "Parents seeking practical readiness skills",
            "Families who want shared awareness training",
            "Couples preparing for greater household readiness",
            "Individuals responsible for the safety of others",
            "Households preparing for greater everyday readiness",
        ],
        "whatYouGet": [
            "Private family-focused instruction",
            "Practical readiness guidance for everyday life",
            "Awareness and emergency-thinking coaching",
            "Customized discussion and training themes",
            "An accessible format for multiple family members",
        ],
        "trainingFocus": [
            "Practical awareness in daily life", "Protective habits and routines",
            "Emergency thinking and response", "Everyday readiness at home and in public",
            "Home-to-street awareness transitions", "Family communication under stress",
        ],
        "likelyOutcomes": [
            "Better shared awareness as a household",
            "Stronger everyday protective habits",
            "More confidence as a family unit",
            "Clearer emergency thinking and response planning",
            "A practical framework families can use in real life",
        ],
        "whyChoose": "This creates a more useful and relevant experience for people who are not just thinking about themselves, but about the people they are responsible for protecting.",
        "faqs": [
            ("Can multiple family members attend?", "Yes. Sessions can include multiple family members when appropriate."),
            ("Is this appropriate for parents and teens?", "Yes. The content is adjusted based on the age and experience of participants."),
            ("Is this more discussion-based or physical?", "It is a blend. Sessions include practical coaching, discussion, and hands-on awareness exercises."),
            ("Can it be customized to our concerns?", "Absolutely. Every session is built around the family's specific needs and priorities."),
            ("Is it beginner-friendly?", "Yes. No prior experience is required for any participant."),
        ],
        "videos": [],
    },
    {
        "title": "Private Workshops",
        "route": "/private-workshops",
        "price": "Starting at $1,500",
        "category": "elite",
        "subheadline": "Private workshops for companies, organizations, leadership teams, women's groups, and select audiences seeking a refined, practical training experience.",
        "overview": "Private Workshops are designed for organizations and select groups that want a more elevated and relevant training experience around awareness, de-escalation, readiness, and personal protection principles.",
        "whoFor": [
            "Companies and corporate teams",
            "Leadership teams and executive groups",
            "Organizations seeking practical training",
            "Women's groups and professional networks",
            "Schools or select institutions",
            "Curated private groups",
        ],
        "whatYouGet": [
            "A private workshop format led by David Furie",
            "Tailored topic selection for your group",
            "Practical instruction for group environments",
            "Premium presentation and founder-led teaching",
            "A more refined alternative to generic safety seminars",
        ],
        "trainingFocus": [
            "Awareness and environmental reading", "De-escalation techniques",
            "Readiness and preparedness", "Personal protection principles",
            "Group safety culture", "Practical response thinking",
        ],
        "likelyOutcomes": [
            "More aware teams and group members",
            "Stronger shared language around readiness",
            "More confidence in public or workplace environments",
            "A memorable and practical private training experience",
        ],
        "whyChoose": "These workshops are designed for groups that want something more serious, more customized, and more practical than a standard generic safety talk.",
        "faqs": [
            ("Who are these workshops best for?", "Private Workshops are ideal for companies, leadership teams, women's groups, and organizations that value serious, practical training."),
            ("Can the content be customized?", "Yes. Every workshop is tailored to the group's needs, industry, and specific concerns."),
            ("Are these held onsite or privately arranged?", "Workshops can be arranged at a private location, your office, or another suitable venue."),
            ("What size groups work best?", "Workshops typically work well with groups of 5 to 30 participants, though larger groups can be accommodated."),
            ("Can this be tailored for women's groups or leadership teams?", "Absolutely. Many workshops are specifically designed for women's groups, executive teams, or professional organizations."),
        ],
        "videos": [],
    },
    {
        "title": "Self Defense",
        "route": "/self-defense-in-brooklyn",
        "price": "",
        "category": "sport",
        "subheadline": "Private self defense training in Brooklyn built around awareness, prevention, and decisive real-world response.",
        "overview": "Self Defense training at Fury Combat Systems is built for real-world readiness. Clients learn practical awareness, prevention, hand-to-hand defense, decisive response, and the mental control needed to act under pressure. The goal is not fear-based training. The goal is confidence, preparation, and the ability to protect yourself and the people who matter to you.",
        "whoFor": [
            "Adults seeking practical self defense in Brooklyn",
            "Commuters and city residents focused on safety",
            "Beginners with no prior training",
            "Women, parents, and professionals",
            "Anyone wanting confidence and personal readiness",
        ],
        "whatYouGet": [
            "Private one-on-one instruction with David Furie",
            "Awareness, prevention, and response coaching",
            "Hands-on defensive skill development",
            "Training scaled to your fitness and experience level",
            "Practical tools you can use in real situations",
        ],
        "trainingFocus": [
            "Situational awareness", "Threat recognition", "Defensive movement",
            "Hand-to-hand protection", "Using natural body weapons",
            "Everyday-object awareness", "Fear control", "Confidence under pressure",
        ],
        "likelyOutcomes": [
            "Stronger awareness in everyday environments",
            "More confidence moving through Brooklyn and the city",
            "Practical defensive skill you can actually use",
            "Better composure under pressure",
            "A clearer mental framework for personal safety",
        ],
        "whyChoose": "This is real self defense built around awareness, prevention, and practical response, taught privately by Grandmaster Dr. David Furie. It is not a generic gym class and it is not fear-based training.",
        "faqs": [
            ("Is this beginner-friendly?", "Yes. Every session is shaped around the client's level, with no prior experience required."),
            ("Is self defense useful if I am not athletic?", "Yes. Self defense at Fury Combat is built around awareness, technique, and decision-making, not raw athleticism."),
            ("Can this focus on street safety in Brooklyn?", "Absolutely. Sessions can be tailored to commuting, neighborhood awareness, and real-world city situations."),
            ("Do you teach awareness and prevention?", "Yes. Awareness and prevention are core to every self defense session."),
            ("Is this private instruction only?", "Yes. All self defense training at Fury Combat is private and tailored to the individual."),
        ],
        "videos": [("Fury Combat NYC: Self Defense & Combat", "https://www.youtube.com/embed/lc5-Rp6ZCZE", "Grandmaster Dr. David Furie demonstrates the Fury Combat approach to self defense and real-world combat application.")],
    },
    {
        "title": "Jujitsu",
        "route": "/jujitsu-in-brooklyn",
        "price": "",
        "category": "sport",
        "subheadline": "Private Jujitsu training in Brooklyn focused on leverage, control, and practical close-combat skill.",
        "overview": "Jujitsu training at Fury Combat Systems focuses on leverage, positioning, control, and practical close-combat skill. It is designed for clients who want to understand how technique can overcome size, strength, and pressure. Training may include stand-up control, takedown awareness, grappling concepts, balance disruption, and self-defense application.",
        "whoFor": [
            "Adults learning Jujitsu for self-defense",
            "Martial artists adding grappling skill",
            "Clients who want technique over brute strength",
            "Beginners seeking private instruction",
            "Experienced practitioners refining their game",
        ],
        "whatYouGet": [
            "Private Jujitsu instruction with David Furie",
            "Step-by-step technical development",
            "Drilling and live application work",
            "Grappling concepts adapted for real-world defense",
            "Sessions tailored to your level and goals",
        ],
        "trainingFocus": [
            "Leverage and control", "Close-combat positioning", "Balance disruption",
            "Defensive grappling", "Escapes and counters", "Ground awareness",
            "Practical self-defense application", "Mental discipline",
        ],
        "likelyOutcomes": [
            "Stronger control in close-range situations",
            "Better understanding of leverage and positioning",
            "Improved composure under physical pressure",
            "Practical grappling skill that applies outside the gym",
            "Greater confidence in your defensive ability",
        ],
        "whyChoose": "Private Jujitsu instruction with Grandmaster Dr. David Furie focuses on real, practical skill, not sport-only mechanics. You learn at your own pace with direct feedback every session.",
        "faqs": [
            ("Do I need prior grappling experience?", "No. Sessions are tailored to your level, from complete beginner through advanced."),
            ("Is this sport Jujitsu or self-defense Jujitsu?", "The focus is practical close-combat skill and self-defense application, not sport competition."),
            ("Can technique really overcome a larger attacker?", "Yes. Jujitsu is built on leverage, positioning, and timing, which is exactly why it works against stronger opponents."),
            ("Is this private instruction only?", "Yes. All Jujitsu training at Fury Combat is private and one-on-one."),
            ("Can this be combined with striking training?", "Absolutely. Many clients combine Jujitsu sessions with kickboxing or MMA training."),
        ],
        "videos": [("Fury Combat NYC: Jujitsu & Self Defense", "https://www.youtube.com/embed/Zu_CTxnVnT0", "A look at leverage, control, and practical close-combat skill from the Fury Combat Jujitsu program.")],
    },
    {
        "title": "Ninjutsu",
        "route": "/ninjutsu-in-brooklyn",
        "price": "",
        "category": "sport",
        "subheadline": "Private Ninjutsu training in Brooklyn built around adaptability, awareness, and practical combat strategy.",
        "overview": "Ninjutsu training at Fury Combat Systems is built around adaptability, awareness, movement, and practical combat strategy. The training connects hand-to-hand skill, defensive movement, weapons concepts, and mental discipline into a private instruction format shaped around the client's level and goals.",
        "whoFor": [
            "Adults drawn to traditional combat arts",
            "Martial artists seeking adaptability and movement",
            "Clients interested in weapons concepts",
            "Beginners curious about Ninjutsu",
            "Practitioners looking for serious private instruction",
        ],
        "whatYouGet": [
            "Private Ninjutsu instruction with David Furie",
            "Founder-led teaching across multiple ranges of combat",
            "Movement, awareness, and timing development",
            "Introduction to weapons concepts where appropriate",
            "A serious training experience shaped around your goals",
        ],
        "trainingFocus": [
            "Adaptability", "Stealth and awareness concepts", "Hand-to-hand movement",
            "Defensive tactics", "Weapons familiarity", "Mind-body discipline",
            "Precision and timing", "Real-world application",
        ],
        "likelyOutcomes": [
            "Greater body awareness and control",
            "Improved adaptability under pressure",
            "Sharper timing and movement",
            "Stronger mind-body discipline",
            "A deeper understanding of combat strategy",
        ],
        "whyChoose": "Ninjutsu at Fury Combat is taught privately by Grandmaster Dr. David Furie with a focus on practical adaptability, not theatrical movement. Every session is built around real skill and real readiness.",
        "faqs": [
            ("Is Ninjutsu practical for modern self-defense?", "Yes. The training emphasizes adaptability, awareness, and real-world application."),
            ("Do I need prior martial arts experience?", "No. Sessions are tailored to your level, whether beginner or experienced."),
            ("Are weapons part of the training?", "Weapons familiarity can be introduced where appropriate, based on the client's goals."),
            ("Is this private instruction only?", "Yes. All Ninjutsu training at Fury Combat is private and one-on-one."),
            ("How is this different from other martial arts?", "Ninjutsu emphasizes adaptability, movement, awareness, and the integration of multiple combat ranges."),
        ],
        "videos": [("Ninjutsu in Brooklyn with Grandmaster Dr. David Furie", "https://www.youtube.com/embed/u2DZiwY8WQA", "Inside the Fury Combat Ninjutsu Shadow Warriors program: adaptability, awareness, and practical combat strategy.")],
    },
    {
        "title": "Kickboxing",
        "route": "/kickboxing-in-brooklyn",
        "price": "",
        "category": "sport",
        "subheadline": "Private kickboxing training in Brooklyn for striking skill, conditioning, and defensive readiness.",
        "overview": "Kickboxing training at Fury Combat Systems develops powerful striking, movement, conditioning, and practical defensive skill. Clients work on punches, kicks, footwork, guard, balance, timing, and reaction in a private training format that can support fitness, self-defense, or mixed martial arts development.",
        "whoFor": [
            "Adults learning kickboxing for fitness or self-defense",
            "Beginners wanting private striking instruction",
            "Experienced fighters refining technique",
            "MMA-focused clients building striking skill",
            "Anyone seeking serious, structured striking training",
        ],
        "whatYouGet": [
            "Private kickboxing instruction with David Furie",
            "Technical breakdown of punches and kicks",
            "Pad work, footwork, and reaction drills",
            "Conditioning built into the session",
            "Defensive striking and guard development",
        ],
        "trainingFocus": [
            "Punching and kicking mechanics", "Footwork", "Guard and defensive awareness",
            "Balance and coordination", "Pad work", "Reaction drills",
            "Conditioning", "Self-defense striking",
        ],
        "likelyOutcomes": [
            "Sharper, more powerful striking",
            "Improved footwork and movement",
            "Better conditioning and stamina",
            "Stronger defensive habits",
            "Practical striking skill that translates to self-defense",
        ],
        "whyChoose": "Private kickboxing instruction with Grandmaster Dr. David Furie means direct feedback every round, technique built correctly from the start, and a session shaped around your goals.",
        "faqs": [
            ("Is this beginner-friendly?", "Yes. Every session is scaled to your fitness level and experience."),
            ("Will I do pad work and live drilling?", "Yes. Pad work, drilling, and reaction work are core parts of the training."),
            ("Is this for fitness or fighting?", "It can be either. Sessions can focus on fitness, self-defense, or MMA development based on your goals."),
            ("Do I need my own gear?", "Some basic gear is helpful, but David will guide you on what you actually need."),
            ("Is this private instruction only?", "Yes. All kickboxing training at Fury Combat is private and one-on-one."),
        ],
        "videos": [("Fury Combat NYC: Kickboxing & MMA", "https://www.youtube.com/embed/EU5wL_VsvEY", "Striking skill, footwork, and conditioning from the Fury Combat kickboxing and MMA program.")],
    },
    {
        "title": "Mixed Martial Arts",
        "route": "/mixed-martial-arts-in-brooklyn",
        "price": "",
        "category": "sport",
        "subheadline": "Private MMA training in Brooklyn combining striking, grappling, and full-spectrum combat skill.",
        "overview": "Mixed Martial Arts training at Fury Combat Systems combines striking, grappling, clinch awareness, ground control, and practical fighting strategy. Clients learn how different ranges of combat connect, including stand-up fighting, defensive movement, close-range control, and ground awareness.",
        "whoFor": [
            "Adults learning MMA in Brooklyn",
            "Fighters building a complete skill set",
            "Beginners wanting structured private MMA instruction",
            "Martial artists cross-training across disciplines",
            "Clients seeking full-range combat readiness",
        ],
        "whatYouGet": [
            "Private MMA instruction with David Furie",
            "Training across striking, clinch, and grappling",
            "Drilling and applied scenario work",
            "Conditioning and movement built into each session",
            "Sessions tailored to your goals and experience",
        ],
        "trainingFocus": [
            "Striking", "Grappling fundamentals", "Clinch fighting", "Ground awareness",
            "Submission concepts", "Sprawl-and-brawl strategy",
            "Defensive positioning", "Full-spectrum combat readiness",
        ],
        "likelyOutcomes": [
            "A more complete fighting skill set",
            "Comfort across striking, clinch, and ground ranges",
            "Improved conditioning and movement",
            "Stronger defensive instincts in every range",
            "Greater confidence in your overall combat ability",
        ],
        "whyChoose": "Private MMA training with Grandmaster Dr. David Furie connects every range of combat into one coherent system, taught directly with no distractions and no wasted time.",
        "faqs": [
            ("Do I need prior martial arts experience?", "No. Sessions are tailored to your level, from complete beginner through advanced."),
            ("Is this for competition or self-defense?", "Either. Sessions can focus on competition preparation, self-defense, or general MMA skill."),
            ("Will I train both striking and grappling?", "Yes. MMA at Fury Combat covers striking, clinch, and ground in an integrated way."),
            ("Is this private instruction only?", "Yes. All MMA training at Fury Combat is private and one-on-one."),
            ("Can I focus on just one area?", "Yes. Sessions can specialize in striking, grappling, or any specific area you want to develop."),
        ],
        "videos": [("Fury Combat NYC: Tactics and Defense Training, Episode 1", "https://www.youtube.com/embed/y-MLTEh0INI", "An inside look at full-spectrum combat training combining striking, grappling, and real-world tactics.")],
    },
    {
        "title": "Weapons and Tactics",
        "route": "/weapons-and-tactics-in-brooklyn",
        "price": "",
        "category": "elite",
        "subheadline": "Private weapons and tactics training in Brooklyn focused on awareness, control, and disciplined response.",
        "overview": "Weapons and Tactics training at Fury Combat Systems introduces clients to practical awareness, tactical movement, weapons familiarity, and disciplined response. The focus is controlled, responsible instruction designed to build readiness, coordination, confidence, and a deeper understanding of real-world defensive situations.",
        "whoFor": [
            "Adults seeking weapons familiarity and tactical awareness",
            "Martial artists expanding their training",
            "Clients focused on serious personal readiness",
            "Beginners wanting responsible, disciplined instruction",
            "Experienced practitioners refining tactical skill",
        ],
        "whatYouGet": [
            "Private weapons and tactics instruction with David Furie",
            "Responsible, controlled introduction to weapons concepts",
            "Tactical movement and awareness drills",
            "Coordination and timing development",
            "A serious training experience shaped around your goals",
        ],
        "trainingFocus": [
            "Weapons familiarity", "Tactical movement", "Awareness and distance",
            "Responsible control", "Defensive positioning", "Improvised-object awareness",
            "Coordination and timing", "Practical readiness",
        ],
        "likelyOutcomes": [
            "Greater tactical awareness in real environments",
            "Stronger discipline and control",
            "Improved coordination and timing",
            "Better understanding of defensive movement",
            "A more complete personal readiness skill set",
        ],
        "whyChoose": "Weapons and Tactics training at Fury Combat is serious, controlled, and taught privately by Grandmaster Dr. David Furie. The focus is responsibility, awareness, and real skill, not flash.",
        "faqs": [
            ("Is this safe for beginners?", "Yes. Training is responsible, controlled, and scaled to the client's level."),
            ("What kinds of weapons concepts are covered?", "Coverage depends on the client's goals and experience, with an emphasis on awareness, movement, and responsible handling."),
            ("Is this only for advanced practitioners?", "No. Sessions are tailored to beginners through advanced clients."),
            ("Is this private instruction only?", "Yes. All Weapons and Tactics training is private and one-on-one."),
            ("Can this be combined with other training?", "Yes. Many clients combine Weapons and Tactics with Ninjutsu, MMA, or private instruction sessions."),
        ],
        "videos": [("Fury Combat NYC: Combat Martial Arts, Self Defense & Survival Tactics", "https://www.youtube.com/embed/vdCp0DTsQok", "Weapons familiarity, tactical movement, and disciplined response from the Fury Combat survival tactics program.")],
    },
]


doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.7)
    s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.8)
    s.right_margin = Inches(0.8)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)


def sp(p, before=0, after=4):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)


def add_run(p, text, size=10.5, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r


# Cover / title
p = doc.add_paragraph(); sp(p, 0, 2); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "FURY COMBAT SYSTEMS", 26, bold=True, color=DARK)
p = doc.add_paragraph(); sp(p, 0, 2); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Complete Service Guide", 13, bold=True, color=RED)
p = doc.add_paragraph(); sp(p, 0, 2); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Private, founder-led training with Soke David Fury (Grandmaster Dr. David Furie)", 10, italic=True, color=GREY)
p = doc.add_paragraph(); sp(p, 0, 10); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Brooklyn, NY  |  (917) 340-2911  |  david.furie@gmail.com", 10, italic=True, color=GREY)


def bullets(items):
    for it in items:
        bp = doc.add_paragraph(style="List Bullet")
        sp(bp, 0, 1)
        add_run(bp, it, 10)


def label(text):
    lp = doc.add_paragraph(); sp(lp, 6, 2)
    add_run(lp, text.upper(), 9.5, bold=True, color=RED)


order = ["elite", "sport"]
svc_num = 0
for cat in order:
    cat_meta = CATEGORIES[cat]
    # Category heading
    hp = doc.add_paragraph(); sp(hp, 10, 1)
    add_run(hp, cat_meta["label"], 16, bold=True, color=DARK)
    tp = doc.add_paragraph(); sp(tp, 0, 1)
    add_run(tp, cat_meta["tagline"], 10.5, bold=True, color=RED)
    dp = doc.add_paragraph(); sp(dp, 0, 6)
    add_run(dp, cat_meta["description"], 10, italic=True, color=GREY)

    for svc in [x for x in SERVICES if x["category"] == cat]:
        svc_num += 1
        # divider line via heading
        title_p = doc.add_paragraph(); sp(title_p, 8, 1)
        add_run(title_p, f"{svc_num}. {svc['title']}", 14, bold=True, color=DARK)

        meta_p = doc.add_paragraph(); sp(meta_p, 0, 2)
        price = svc["price"] if svc["price"] else "Custom pricing (by inquiry)"
        add_run(meta_p, f"Price: ", 10, bold=True, color=RED)
        add_run(meta_p, f"{price}      ", 10)
        add_run(meta_p, f"Page: ", 10, bold=True, color=RED)
        add_run(meta_p, svc["route"], 10)

        sub_p = doc.add_paragraph(); sp(sub_p, 0, 4)
        add_run(sub_p, svc["subheadline"], 10.5, italic=True, color=GREY)

        ov_p = doc.add_paragraph(); sp(ov_p, 0, 4)
        add_run(ov_p, svc["overview"], 10.5)

        label("Who it's for")
        bullets(svc["whoFor"])
        label("What you get")
        bullets(svc["whatYouGet"])
        label("Training focus")
        tf_p = doc.add_paragraph(); sp(tf_p, 0, 2)
        add_run(tf_p, "  •  ".join(svc["trainingFocus"]), 10)
        label("Likely outcomes")
        bullets(svc["likelyOutcomes"])
        label("Why choose")
        wc_p = doc.add_paragraph(); sp(wc_p, 0, 4)
        add_run(wc_p, svc["whyChoose"], 10.5)
        label("Frequently asked questions")
        for q, a in svc["faqs"]:
            qp = doc.add_paragraph(); sp(qp, 2, 0)
            add_run(qp, f"Q: {q}", 10, bold=True, color=DARK)
            ap = doc.add_paragraph(); sp(ap, 0, 2)
            add_run(ap, f"A: {a}", 10)
        if svc["videos"]:
            label("Featured video")
            for vt, vurl, vdesc in svc["videos"]:
                vp = doc.add_paragraph(); sp(vp, 0, 1)
                add_run(vp, vt, 10, bold=True)
                vdp = doc.add_paragraph(); sp(vdp, 0, 1)
                add_run(vdp, vdesc, 10, italic=True, color=GREY)
                vup = doc.add_paragraph(); sp(vup, 0, 3)
                add_run(vup, vurl.replace("/embed/", "/watch?v="), 9.5, color=GREY)

out = "Fury-Combat-Systems-Services.docx"
doc.save(out)
print("saved", out, "with", svc_num, "services")
